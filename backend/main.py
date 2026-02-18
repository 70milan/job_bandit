# backend/main.py
from fastapi import FastAPI, WebSocket, Request
from fastapi.responses import StreamingResponse
from openai import OpenAI
import os
import sys
import io

import json
import asyncio
import websockets
from fastapi import FastAPI, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from pathlib import Path
from typing import Optional, Dict, Any

# Token counting for cost estimation
try:
    import tiktoken
    encoding = tiktoken.encoding_for_model("gpt-4o")
except:
    tiktoken = None
    encoding = None
    print("Warning: tiktoken not installed - cost estimation will be approximate")

# Session usage tracking
session_usage = {
    "input_tokens": 0,
    "output_tokens": 0,
    "total_cost": 0.0,
    "request_count": 0
}

# GPT pricing (per 1M tokens) - Updated December 2025
PRICING = {
    "gpt-3.5-turbo": {"input": 0.50, "output": 1.50},
    "gpt-4o-mini": {"input": 0.60, "output": 2.40},
    "gpt-4o": {"input": 2.50, "output": 10.00},
    "gpt-5-nano": {"input": 0.30, "output": 1.20},
    "gpt-5-mini": {"input": 1.50, "output": 6.00},
    "gpt-5": {"input": 5.00, "output": 20.00}
}

# Available models for text-only responses (user can select)
AVAILABLE_TEXT_MODELS = {
    "gpt-3.5-turbo": {"name": "GPT-3.5 Turbo", "speed": "Fastest", "cost": "Cheapest", "accuracy": "Good", "desc": "Quick text answers"},
    "gpt-4o-mini": {"name": "GPT-4o Mini", "speed": "Fast", "cost": "Cheap", "accuracy": "Better", "desc": "Balanced option"},
    "gpt-4o": {"name": "GPT-4o", "speed": "Medium", "cost": "Moderate", "accuracy": "Great", "desc": "Complex questions"},
    "gpt-5-nano": {"name": "GPT-5 Nano", "speed": "Fast", "cost": "Cheap", "accuracy": "Good", "desc": "Lightweight GPT-5"},
    "gpt-5-mini": {"name": "GPT-5 Mini", "speed": "Medium", "cost": "Moderate", "accuracy": "Better", "desc": "Balanced GPT-5"},
    "gpt-5": {"name": "GPT-5", "speed": "Slower", "cost": "Premium", "accuracy": "Best", "desc": "Maximum quality"}
}

DEFAULT_TEXT_MODEL = "gpt-4o"

def count_tokens(text: str) -> int:
    """Count tokens in text"""
    if encoding:
        return len(encoding.encode(text))
    return len(text) // 4  # Rough estimate

def estimate_image_tokens(base64_data: str, model: str = "gpt-4o") -> int:
    """Estimate tokens for an image based on base64 data size.
    GPT-4o uses ~255 tokens per 512x512 tile.
    We estimate image dimensions from base64 size and calculate tiles.
    """
    try:
        # Remove data URL prefix if present
        if "base64," in base64_data:
            base64_data = base64_data.split("base64,")[1]
        
        # Estimate image size from base64 length
        # Base64 is ~4/3 of original size, JPEG compression is ~10:1
        byte_size = len(base64_data) * 3 // 4
        estimated_pixels = byte_size * 10  # Rough decompression estimate
        
        # Estimate dimensions (assume roughly square)
        import math
        side = int(math.sqrt(estimated_pixels))
        
        # GPT-4o: 512x512 tiles, ~255 tokens per tile + 85 base tokens
        tiles_x = max(1, (side + 511) // 512)
        tiles_y = max(1, (side + 511) // 512)
        total_tiles = tiles_x * tiles_y
        
        # Base 85 tokens + 255 per tile (GPT-4o vision formula)
        image_tokens = 85 + (255 * total_tiles)
        
        print(f"[IMAGE TOKENS] Estimated {image_tokens} tokens for ~{side}x{side} image ({total_tiles} tiles)")
        return image_tokens
    except Exception as e:
        print(f"[IMAGE TOKENS] Error estimating: {e}, using default 500")
        return 500  # Safe default for a medium image

def calculate_cost(input_tokens: int, output_tokens: int, model: str = "gpt-4o", image_tokens: int = 0) -> float:
    """Calculate cost in dollars"""
    pricing = PRICING.get(model, PRICING["gpt-4o"])
    total_input = input_tokens + image_tokens
    input_cost = (total_input / 1_000_000) * pricing["input"]
    output_cost = (output_tokens / 1_000_000) * pricing["output"]
    return (input_cost + output_cost) * 1.10  # 10% buffer

def update_usage(input_tokens: int, output_tokens: int, model: str = "gpt-3.5-turbo", image_tokens: int = 0):
    """Update session usage stats"""
    global session_usage
    session_usage["input_tokens"] += input_tokens + image_tokens
    session_usage["output_tokens"] += output_tokens
    session_usage["total_cost"] += calculate_cost(input_tokens, output_tokens, model, image_tokens)
    session_usage["request_count"] += 1
    print(f"[USAGE] Total: ${session_usage['total_cost']:.4f} ({session_usage['request_count']} requests, {image_tokens} image tokens)")

# Fix Unicode encoding for Windows console (prevents charmap errors with special characters)
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

app = FastAPI()

# Determine base directory for data storage (persists across updates)
if getattr(sys, 'frozen', False):
    # Running as compiled executable - store data in AppData to survive updates
    # e.g., C:\Users\<user>\AppData\Roaming\JobAndit\backend
    app_data = os.getenv('APPDATA')
    if app_data:
        BASE_DIR = Path(app_data) / "JobAndit" / "backend"
        BASE_DIR.mkdir(parents=True, exist_ok=True)
    else:
        # Fallback if APPDATA not found (rare)
        BASE_DIR = Path(sys.executable).parent
else:
    # Running as script - store locally
    BASE_DIR = Path(__file__).parent

# Allow the Electron frontend and local testing to call the API
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Debug middleware to log raw requests
@app.middleware("http")
async def log_requests(request: Request, call_next):
    if request.url.path == "/ai":
        body = await request.body()
        print(f"\n[RAW REQUEST] /ai received {len(body)} bytes")
        try:
            data = json.loads(body)
            print(f"[RAW REQUEST] Keys: {list(data.keys())}")
            for k, v in data.items():
                if k == 'screenshot':
                    print(f"  {k}: <{len(str(v))} chars>")
                elif isinstance(v, str) and len(v) > 100:
                    print(f"  {k}: {v[:100]}...")
                else:
                    print(f"  {k}: {v}")
        except Exception as e:
            print(f"[RAW REQUEST] Failed to parse JSON: {e}")
            print(f"[RAW REQUEST] Body preview: {body[:500]}")
    response = await call_next(request)
    return response

def get_api_key():
    """Get API key from user profile only"""
    global profile_cache
    if isinstance(profile_cache, dict) and profile_cache.get('openai_api_key'):
        return profile_cache['openai_api_key']
    return None

@app.websocket("/realtime")
async def realtime(ws: WebSocket):
    await ws.accept()
    
    api_key = get_api_key()
    url = "wss://api.openai.com/v1/realtime?model=gpt-4o-realtime-preview-2024-10-01"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "OpenAI-Beta": "realtime=v1",
    }
    
    print(f"Incoming WebSocket connection. API Key present: {bool(api_key)}")
    if not api_key:
        print("ERROR: API key not found in profile!")
        await ws.close(code=1008, reason="Missing API Key")
        return

    print(f"Connecting to OpenAI at {url}...")
    try:
        async with websockets.connect(url, additional_headers=headers) as openai_ws:
            print("Connected to OpenAI Realtime API!")
            # Initialize session
            session_update = {
                "type": "session.update",
                "session": {
                    "modalities": ["text"],
                    "input_audio_transcription": {
                        "model": "whisper-1"
                    }
                }
            }
            await openai_ws.send(json.dumps(session_update))

            async def receive_from_client():
                try:
                    while True:
                        data = await ws.receive_text()
                        # Expecting base64 audio from client
                        # Send to OpenAI
                        event = {
                            "type": "input_audio_buffer.append",
                            "audio": data
                        }
                        await openai_ws.send(json.dumps(event))
                except WebSocketDisconnect:
                    pass
                except Exception as e:
                    print(f"Client receive error: {e}")

            async def receive_from_openai():
                try:
                    async for message in openai_ws:
                        event = json.loads(message)
                        
                        # Real-time transcription events
                        if event["type"] == "conversation.item.input_audio_transcription.delta":
                            print(f"TRANSCRIPT DELTA: {event['delta']}")
                            await ws.send_json({
                                "type": "transcript",
                                "text": event["delta"]
                            })
                        elif event["type"] == "conversation.item.input_audio_transcription.completed":
                            print(f"TRANSCRIPT DONE: {event['transcript']}")
                            await ws.send_json({
                                "type": "transcript",
                                "text": "\n"
                            })
                        elif event["type"] == "error":
                            print(f"OpenAI Error: {event}")
                except Exception as e:
                    print(f"OpenAI receive error: {e}")

            await asyncio.gather(receive_from_client(), receive_from_openai())

    except websockets.exceptions.ConnectionClosed as e:
        print(f"OpenAI Connection Closed: {e.code} {e.reason}")
        await ws.close(code=1011, reason=f"OpenAI Closed: {e.code}")
    except Exception as e:
        print(f"Connection error: {e}")
        # Try to send the error to the client before closing
        try:
            await ws.close(code=1011, reason=f"Upstream Error: {str(e)[:100]}")
        except:
            pass

from pydantic import BaseModel
from fastapi import File, UploadFile, Form
try:
    from docx import Document
except ModuleNotFoundError:
    Document = None
    print("Warning: python-docx not installed — resume upload endpoint will be unavailable until installed.")


class AIRequest(BaseModel):
    transcript: str = ""
    role: str = "data engineer"
    screenshot: Optional[str] = None
    job_description: Optional[str] = None
    save_to_context: Optional[bool] = True  # Set False for one-shot problems (LeetCode), True for scenarios needing follow-up
    text_model: Optional[str] = None  # Selected model for text-only responses

    class Config:
        extra = "ignore"  # Ignore extra fields


@app.get('/models')
async def get_available_models():
    """Return available text models with metadata for UI display"""
    models = []
    for model_id, info in AVAILABLE_TEXT_MODELS.items():
        pricing = PRICING.get(model_id, {})
        models.append({
            "id": model_id,
            "name": info["name"],
            "speed": info["speed"],
            "cost": info["cost"],
            "accuracy": info["accuracy"],
            "description": info["desc"],
            "price_input": pricing.get("input", 0),
            "price_output": pricing.get("output", 0)
        })
    return {"models": models, "default": DEFAULT_TEXT_MODEL}


# Debug endpoint to test raw requests
@app.post("/ai/debug")
async def debug_ai_request(request: Dict[str, Any]):
    print(f"\n[DEBUG RAW] Received keys: {list(request.keys())}")
    for k, v in request.items():
        if k == 'screenshot':
            print(f"  {k}: <{len(str(v))} chars>")
        else:
            print(f"  {k}: {v}")
    return {"status": "ok", "received_keys": list(request.keys())}


# Persistent profile support -------------------------------------------------
PROFILE_PATH = BASE_DIR / "user_profile.json"

def load_profile() -> Dict[str, Any]:
    if PROFILE_PATH.exists():
        try:
            return json.loads(PROFILE_PATH.read_text(encoding="utf-8"))
        except Exception as e:
            print(f"Error loading profile: {e}")
            return {}
    return {}

def save_profile(data: Dict[str, Any]):
    try:
        json_str = json.dumps(data, indent=2, ensure_ascii=False)
        PROFILE_PATH.write_text(json_str, encoding="utf-8")
        print(f"[SAVE_PROFILE] Successfully wrote {len(json_str)} bytes to {PROFILE_PATH}")
    except Exception as e:
        print(f"[SAVE_PROFILE ERROR] Failed to save: {e}")
        raise  # Re-raise so caller knows it failed
        print(f"Error saving profile: {e}")

profile_cache = load_profile()

# Conversation history for session context
# Stores list of {"role": "user"/"assistant", "content": "..."} messages
conversation_history = []

# Auto-reset disabled - resume will persist between restarts
# Uncomment below to enable auto-clearing on startup:
# if profile_cache and isinstance(profile_cache, dict):
#     if 'resume_text' in profile_cache and profile_cache['resume_text']:
#         profile_cache['resume_text'] = ""
#         save_profile(profile_cache)
#         print("[STARTUP] Resume text cleared - ready for new upload")
#     else:
#         print("[STARTUP] No resume to clear - ready for upload")

@app.get('/profile')
async def get_profile():
    return profile_cache

@app.post('/validate-api-key')
async def validate_api_key(data: Dict[str, str]):
    """Validate OpenAI API key before allowing session creation"""
    api_key = data.get('api_key', '').strip()
    
    if not api_key:
        return {"valid": False, "error": "API key is required"}
    
    if not api_key.startswith('sk-'):
        return {"valid": False, "error": "Invalid API key format"}
    
    try:
        client = OpenAI(api_key=api_key)
        client.models.list()
        return {"valid": True}
    except Exception as e:
        error_msg = str(e)
        if "invalid_api_key" in error_msg or "Incorrect API key" in error_msg:
            return {"valid": False, "error": "Invalid API key"}
        elif "insufficient_quota" in error_msg:
            return {"valid": False, "error": "API key has no credits/quota"}
        else:
            return {"valid": False, "error": f"API key validation failed: {error_msg[:100]}"}

# License validation - key stored securely on backend
VALID_LICENSE_KEYS = set()  # Scrubbed from history

@app.post('/validate-license')
async def validate_license(data: Dict[str, str]):
    """Validate license key for full session access"""
    license_key = data.get('license_key', '').strip()
    
    if not license_key:
        return {"valid": False, "status": "empty"}
    
    if license_key in VALID_LICENSE_KEYS:
        return {"valid": True, "status": "valid"}
    
    return {"valid": False, "status": "invalid"}

@app.post('/profile')
async def post_profile(data: Dict[str, Any]):
    global profile_cache
    try:
        save_profile(data)
        profile_cache = data
        return {"status": "ok"}
    except Exception as e:
        return {"status": "error", "error": str(e)}


# ============ SESSION MANAGEMENT ============
# Sessions are stored in: BASE_DIR/sessions/<session_name>/
# Each session folder contains:
#   - session.json (metadata, job description, resume text)
#   - resume.docx (original resume file)
#   - conversation.json (all Q&A pairs)

SESSIONS_DIR = BASE_DIR / 'sessions'
current_session_name = None

def save_conversation_to_session(question: str, response: str, had_screenshot: bool = False):
    """Helper function to save a Q&A pair to the current session"""
    global current_session_name
    
    if not current_session_name:
        return
    
    session_dir = SESSIONS_DIR / current_session_name
    conv_file = session_dir / 'conversation.json'
    
    try:
        if conv_file.exists():
            conversation = json.loads(conv_file.read_text(encoding='utf-8'))
        else:
            conversation = []
        
        from datetime import datetime
        entry = {
            'timestamp': datetime.now().isoformat(),
            'question': question,
            'response': response,
            'had_screenshot': had_screenshot
        }
        conversation.append(entry)
        
        conv_file.write_text(json.dumps(conversation, indent=2), encoding='utf-8')
        print(f"[SESSION] Auto-saved conversation entry #{len(conversation)}")
    except Exception as e:
        print(f"[SESSION SAVE ERROR] {e}")

@app.post('/session/create')
async def create_session(data: Dict[str, str]):
    """Create a new session folder"""
    global current_session_name
    
    session_name = data.get('session_name', '').strip()
    if not session_name:
        return {"status": "error", "error": "Session name is required"}
    
    session_dir = SESSIONS_DIR / session_name
    try:
        session_dir.mkdir(parents=True, exist_ok=True)
        current_session_name = session_name
        print(f"[SESSION] Created session folder: {session_dir}")
        return {"status": "ok", "session_path": str(session_dir)}
    except Exception as e:
        return {"status": "error", "error": str(e)}


@app.post('/session/resume')
async def upload_session_resume(file: UploadFile = File(...), session_name: str = Form(...)):
    """Upload resume to a specific session folder"""
    global profile_cache, current_session_name
    
    try:
        filename = file.filename
        if not filename.lower().endswith('.docx'):
            return {"status": "error", "error": "Only .docx files supported"}
        
        session_dir = SESSIONS_DIR / session_name
        session_dir.mkdir(parents=True, exist_ok=True)
        
        # Save resume file to session folder
        resume_path = session_dir / filename
        content = await file.read()
        resume_path.write_bytes(content)
        print(f"[SESSION] Resume saved to: {resume_path}")
        
        # Extract text from docx
        from docx import Document
        from io import BytesIO
        doc = Document(BytesIO(content))
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        current_session_name = session_name
        
        # Also update profile cache for AI responses
        if profile_cache is None:
            profile_cache = {}
        profile_cache['resume_text'] = text
        
        return {"status": "ok", "resume_text": text, "resume_path": str(resume_path)}
    except Exception as e:
        print(f"[SESSION RESUME ERROR] {e}")
        return {"status": "error", "error": str(e)}


@app.post('/session/save')
async def save_session_data(data: Dict[str, Any]):
    """Save session metadata (job description, API key, etc.)"""
    global profile_cache, current_session_name
    
    session_name = data.get('session_name', '').strip()
    if not session_name:
        return {"status": "error", "error": "Session name is required"}
    
    session_dir = SESSIONS_DIR / session_name
    session_dir.mkdir(parents=True, exist_ok=True)
    
    session_file = session_dir / 'session.json'
    
    try:
        # Save session data
        session_data = {
            'session_name': session_name,
            'job_description': data.get('job_description', ''),
            'resume_text': data.get('resume_text', ''),
            'created_at': data.get('created_at', ''),
            'updated_at': str(Path('').resolve()),  # Will be updated on each save
            'text_model': data.get('text_model', DEFAULT_TEXT_MODEL)
        }
        
        session_file.write_text(json.dumps(session_data, indent=2), encoding='utf-8')
        print(f"[SESSION] Saved session data to: {session_file}")
        
        # Initialize empty conversation file
        conv_file = session_dir / 'conversation.json'
        if not conv_file.exists():
            conv_file.write_text('[]', encoding='utf-8')
        
        # Update profile cache for AI
        if profile_cache is None:
            profile_cache = {}
        profile_cache['openai_api_key'] = data.get('openai_api_key', '')
        profile_cache['job_description'] = data.get('job_description', '')
        profile_cache['resume_text'] = data.get('resume_text', '')
        
        # Also save to profile for persistence
        save_profile(profile_cache)
        
        current_session_name = session_name
        
        return {"status": "ok", "session_path": str(session_dir)}
    except Exception as e:
        print(f"[SESSION SAVE ERROR] {e}")
        return {"status": "error", "error": str(e)}


@app.post('/session/conversation')
async def save_conversation_entry(data: Dict[str, Any]):
    """Save a conversation entry (question + response) to the session"""
    global current_session_name
    
    session_name = data.get('session_name') or current_session_name
    if not session_name:
        return {"status": "error", "error": "No active session"}
    
    session_dir = SESSIONS_DIR / session_name
    conv_file = session_dir / 'conversation.json'
    
    try:
        # Load existing conversation
        if conv_file.exists():
            conversation = json.loads(conv_file.read_text(encoding='utf-8'))
        else:
            conversation = []
        
        # Add new entry
        entry = {
            'timestamp': data.get('timestamp', ''),
            'question': data.get('question', ''),
            'response': data.get('response', ''),
            'had_screenshot': data.get('had_screenshot', False)
        }
        conversation.append(entry)
        
        # Save back
        conv_file.write_text(json.dumps(conversation, indent=2), encoding='utf-8')
        print(f"[SESSION] Saved conversation entry #{len(conversation)} to: {conv_file}")
        
        return {"status": "ok", "entry_count": len(conversation)}
    except Exception as e:
        print(f"[SESSION CONVERSATION ERROR] {e}")
        return {"status": "error", "error": str(e)}


@app.post('/session/end')
async def end_session(data: Dict[str, str]):
    """Finalize and close a session"""
    global current_session_name, conversation_history
    
    session_name = data.get('session_name') or current_session_name
    if not session_name:
        return {"status": "ok", "message": "No session to end"}
    
    session_dir = SESSIONS_DIR / session_name
    
    try:
        # Save final conversation history to session
        if conversation_history:
            conv_file = session_dir / 'conversation.json'
            if conv_file.exists():
                existing = json.loads(conv_file.read_text(encoding='utf-8'))
            else:
                existing = []
            
            # Add any remaining history
            for i in range(0, len(conversation_history), 2):
                if i + 1 < len(conversation_history):
                    entry = {
                        'timestamp': '',
                        'question': conversation_history[i].get('content', ''),
                        'response': conversation_history[i + 1].get('content', ''),
                        'had_screenshot': False
                    }
                    existing.append(entry)
            
            conv_file.write_text(json.dumps(existing, indent=2), encoding='utf-8')
        
        # Clear current session
        current_session_name = None
        conversation_history = []
        
        print(f"[SESSION] Ended session: {session_name}")
        return {"status": "ok", "message": f"Session '{session_name}' ended"}
    except Exception as e:
        return {"status": "error", "error": str(e)}


@app.get('/sessions')
async def list_sessions():
    """List all available sessions"""
    try:
        if not SESSIONS_DIR.exists():
            return {"sessions": []}
        
        sessions = []
        for session_dir in SESSIONS_DIR.iterdir():
            if session_dir.is_dir():
                session_file = session_dir / 'session.json'
                if session_file.exists():
                    data = json.loads(session_file.read_text(encoding='utf-8'))
                    sessions.append({
                        'name': session_dir.name,
                        'created_at': data.get('created_at', ''),
                        'job_description_preview': data.get('job_description', '')[:100]
                    })
        
        return {"sessions": sessions}
    except Exception as e:
        return {"sessions": [], "error": str(e)}


@app.get('/session/load/{session_name}')
async def load_session(session_name: str):
    """Load a previous session's data"""
    try:
        session_dir = SESSIONS_DIR / session_name
        session_file = session_dir / 'session.json'
        
        if not session_file.exists():
            return {"status": "error", "error": "Session not found"}
        
        data = json.loads(session_file.read_text(encoding='utf-8'))
        
        return {
            "status": "ok",
            "session_name": session_name,
            "resume_text": data.get('resume_text', ''),
            "job_description": data.get('job_description', ''),
            "created_at": data.get('created_at', ''),
            "text_model": data.get('text_model', DEFAULT_TEXT_MODEL)
        }
    except Exception as e:
        return {"status": "error", "error": str(e)}


@app.delete('/session/delete/{session_name}')
async def delete_session(session_name: str):
    """Delete a session and all its contents"""
    try:
        import shutil
        session_dir = SESSIONS_DIR / session_name
        
        if not session_dir.exists():
            return {"status": "error", "error": "Session not found"}
        
        shutil.rmtree(session_dir)
        print(f"[SESSION] Deleted session: {session_name}")
        
        return {"status": "ok", "message": f"Session '{session_name}' deleted"}
    except Exception as e:
        print(f"[SESSION DELETE ERROR] {e}")
        return {"status": "error", "error": str(e)}

# ============ END SESSION MANAGEMENT ============


@app.post('/conversation/clear')
async def clear_conversation():
    """Clear conversation history to start a fresh session"""
    global conversation_history
    conversation_history = []
    print("[CONVERSATION] History cleared - starting fresh session")
    return {"status": "ok", "message": "Conversation history cleared"}


@app.get('/conversation/history')
async def get_conversation_history():
    """Get current conversation history (for debugging)"""
    return {"history": conversation_history, "count": len(conversation_history)}


@app.get('/usage')
async def get_usage():
    """Get current session API usage and cost"""
    return session_usage


@app.post('/usage/reset')
async def reset_usage():
    """Reset usage counters"""
    global session_usage
    session_usage = {
        "input_tokens": 0,
        "output_tokens": 0,
        "total_cost": 0.0,
        "request_count": 0
    }
    return {"status": "ok", "message": "Usage reset"}

@app.post('/profile/resume')
async def upload_resume(file: UploadFile = File(...)):
    global profile_cache
    try:
        filename = file.filename
        if not filename.lower().endswith('.docx'):
            return {"status": "error", "error": "Only .docx files supported"}

        resumes_dir = BASE_DIR / 'resumes'
        print(f"Creating resumes directory at: {resumes_dir}")
        resumes_dir.mkdir(parents=True, exist_ok=True)
        print(f"Directory created: {resumes_dir.exists()}")
        out_path = resumes_dir / filename
        content = await file.read()
        out_path.write_bytes(content)

        # Extract text using python-docx (if available)
        if Document is None:
            return {"status": "error", "error": "python-docx is not installed on the server. Install with: pip install python-docx"}

        try:
            doc = Document(out_path)
            text = "\n\n".join([p.text for p in doc.paragraphs if p.text and p.text.strip()])
        except Exception as e:
            return {"status": "error", "error": f"docx parse failed: {e}"}

        # Save extracted resume text into profile
        if not isinstance(profile_cache, dict):
            profile_cache = {}
        profile_cache['resume_text'] = text
        
        try:
            save_profile(profile_cache)
            print(f"\n{'='*60}")
            print(f"[RESUME UPLOAD SUCCESS]")
            print(f"File: {filename}")
            print(f"Text extracted: {len(text)} characters")
            print(f"Saved to profile!")
            print(f"{'='*60}\n")
        except Exception as save_error:
            print(f"\n[RESUME SAVE FAILED]: {save_error}\n")
            return {"status": "error", "error": f"Failed to save resume: {save_error}"}

        return {"status": "ok", "resume_snippet": text[:800]}
    except Exception as e:
        print(f"\n[RESUME UPLOAD ERROR]: {e}\n")
        return {"status": "error", "error": str(e)}

# ----------------------------------------------------------------------------

@app.post("/ai/stream")
async def stream_ai_response(req: AIRequest):
    """Streaming endpoint for real-time AI responses using Server-Sent Events"""
    global conversation_history
    
    async def generate_stream():
        global conversation_history
        try:
            client = OpenAI(api_key=get_api_key())
            
            # CRITICAL: Reload profile from disk to get latest resume
            current_profile = load_profile()
            
            # Extract resume and job description
            resume_text = None
            job_description = None
            profile_metadata = None
            
            if current_profile and isinstance(current_profile, dict):
                resume_text = current_profile.get('resume_text')
                job_description = current_profile.get('job_description')
                # Build metadata without resume_text
                profile_metadata = {k: v for k, v in current_profile.items() if k not in ['resume_text', 'job_description']}
            
            has_resume = bool(resume_text and resume_text.strip())
            
            # DEBUG: Print what we're sending
            print(f"\n{'='*60}")
            print(f"[STREAMING AI REQUEST]")
            print(f"Has resume: {has_resume}")
            print(f"Conversation history length: {len(conversation_history)} messages")
            if has_resume:
                print(f"Resume length: {len(resume_text)} characters")
            if job_description:
                print(f"Job description length: {len(job_description)} characters")
            print(f"{'='*60}\n")

            if req.screenshot:
                # Use vision model for screenshot analysis
                messages = [
                    {
                        "role": "system",
                        "content": (
                            f"CRITICAL: You ARE the job candidate in this interview for a {req.role} position. "
                            "You are NOT an AI assistant. NEVER say 'I am an AI' or 'I don't have experience'. "
                            "The resume below is YOUR real work history. When asked about past jobs, refer to YOUR resume. "
                            "Speak in FIRST PERSON as yourself - a real human professional. "
                            "When you see a coding problem, SOLVE IT with working code and clear explanation. "
                            "Use SIMPLE, clear words. Avoid complex vocabulary. Speak naturally like in a real conversation. "
                            "Explain your approach in 8-12 sentences. Write clean code with comments, then explain simply. "
                            "Draw from YOUR resume for all experience-related questions. "
                            "Prefer SQL, PySpark, Pandas, Python based on the job. "
                            "Remember what was discussed earlier."
                        )
                    }
                ]

                # Inject resume
                if has_resume:
                    messages.append({"role": "user", "content": "Candidate resume:\n" + resume_text})
                elif profile_metadata:
                    messages.append({"role": "user", "content": "Candidate profile:\n" + json.dumps(profile_metadata, indent=2)})

                # Include job description
                if job_description:
                    messages.append({"role": "user", "content": f"Job description:\n{job_description}"})

                # Add conversation history for context
                messages.extend(conversation_history)

                # Add current question with screenshot
                current_user_msg = {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": req.transcript},
                        {"type": "image_url", "image_url": {"url": req.screenshot}}
                    ]
                }
                messages.append(current_user_msg)
                model = "gpt-4o"
                print(f"[STREAM] Using model: {model} (vision)")
            else:
                # Use text-only model - GPT-3.5-turbo for speed!
                system_content = (
                    f"CRITICAL: You ARE the job candidate in this interview for a {req.role} position. "
                    "You are NOT an AI assistant. NEVER say 'I am an AI' or 'I don't have experience' or 'I cannot'. "
                    "The resume below is YOUR real work history. You ARE this person with this background. "
                    "When asked about previous jobs, projects, or experience - use details from YOUR resume. "
                    "Speak in FIRST PERSON as yourself - a real human professional. "
                    "Use SIMPLE, clear words. Speak naturally like in a real conversation. "
                    "Be confident but humble. Keep answers to 4-6 sentences. Give specific examples from YOUR resume. "
                    "Remember what was discussed earlier in this conversation."
                )

                messages = [
                    {"role": "system", "content": system_content}
                ]

                # Inject resume
                if has_resume:
                    messages.append({"role": "user", "content": "Candidate resume:\n" + resume_text})
                elif profile_metadata:
                    messages.append({"role": "user", "content": "Candidate profile:\n" + json.dumps(profile_metadata, indent=2)})

                # Include job description
                if job_description:
                    messages.append({"role": "user", "content": f"Job description:\n{job_description}"})

                # Add conversation history for context
                messages.extend(conversation_history)

                # Add current question
                messages.append({"role": "user", "content": req.transcript})
                # Use selected model or default
                model = req.text_model if req.text_model and req.text_model in AVAILABLE_TEXT_MODELS else DEFAULT_TEXT_MODEL
                print(f"[STREAM] Using model: {model} (text-only, user selected: {req.text_model})")

            # GPT-5+ models are reasoning models: they use tokens for internal thinking
            # before producing output, so they need a much higher token limit
            token_param = {}
            if model.startswith("gpt-5"):
                token_param["max_completion_tokens"] = 4096
            else:
                token_param["max_tokens"] = 600
                token_param["temperature"] = 0.7
            
            # Send heartbeat to establish SSE connection
            yield f"data: {json.dumps({'heartbeat': True})}\n\n"
            
            # Create completion - catch errors separately so they always reach the client
            try:
                completion = client.chat.completions.create(
                    model=model,
                    messages=messages,
                    stream=True,
                    **token_param
                )
            except Exception as create_err:
                print(f"[STREAM] OpenAI create() error: {create_err}")
                yield f"data: {json.dumps({'error': str(create_err)})}\n\n"
                return
            
            # Collect full response for history
            full_response = ""
            
            # Yield chunks as they arrive - wrap in try/except for iteration errors
            try:
                for chunk in completion:
                    if chunk.choices and chunk.choices[0].delta.content:
                        content = chunk.choices[0].delta.content
                        full_response += content
                        # Send as Server-Sent Event format
                        yield f"data: {json.dumps({'chunk': content})}\n\n"
            except Exception as iter_err:
                err_msg = str(iter_err)[:200]
                print(f"[STREAM] Iteration error: {iter_err}")
                yield f"data: {json.dumps({'error': err_msg})}\n\n"
                return
            
            # Check for empty response
            if not full_response.strip():
                print(f"[STREAM] WARNING: Model {model} returned empty response")
                yield f"data: {json.dumps({'error': f'Model {model} returned an empty response. The model may not support this request format.'})}\n\n"
                return

            # Save to conversation history only if save_to_context is True
            # (skip for one-shot LeetCode problems, save for scenarios needing follow-up)
            if req.save_to_context:
                if req.screenshot:
                    user_msg = f"[USER SHARED A SCREENSHOT] Question about the screenshot: {req.transcript}"
                else:
                    user_msg = req.transcript
                conversation_history.append({"role": "user", "content": user_msg})
                conversation_history.append({"role": "assistant", "content": full_response})
                
                # Save to session folder if active
                if current_session_name:
                    try:
                        save_conversation_to_session(user_msg, full_response, bool(req.screenshot))
                    except Exception as save_err:
                        print(f"[SESSION SAVE ERROR] {save_err}")
                
                # Limit history to last 10 exchanges (20 messages) to control costs
                if len(conversation_history) > 20:
                    conversation_history = conversation_history[-20:]
                
                print(f"[STREAM] Conversation history now has {len(conversation_history)} messages")
            else:
                print(f"[STREAM] Skipped saving to history (save_to_context=False)")
            
            # Calculate and track usage
            input_text = "\n".join([m.get('content', '') if isinstance(m.get('content'), str) else str(m.get('content', '')) for m in messages])
            input_tokens = count_tokens(input_text)
            output_tokens = count_tokens(full_response)
            
            # Estimate image tokens if screenshot was used
            image_tokens = 0
            if req.screenshot:
                image_tokens = estimate_image_tokens(req.screenshot, model)
            
            update_usage(input_tokens, output_tokens, model, image_tokens)
            
            # Send completion signal with usage info
            yield f"data: {json.dumps({'done': True, 'model': model, 'usage': session_usage})}\n\n"
            
        except Exception as e:
            err_msg = str(e)[:200]
            print(f"Streaming AI Error: {e}")
            yield f"data: {json.dumps({'error': err_msg})}\n\n"
    
    return StreamingResponse(generate_stream(), media_type="text/event-stream")


@app.post("/ai")
async def generate_ai_response(req: AIRequest):
    global conversation_history
    print(f"\n[DEBUG] /ai called with: transcript={req.transcript[:50] if req.transcript else None}..., role={req.role}, screenshot={'YES' if req.screenshot else 'NO'}, save_to_context={req.save_to_context}")
    try:
        client = OpenAI(api_key=get_api_key())
        
        # CRITICAL: Reload profile from disk to get latest resume (in case it was uploaded after backend started)
        current_profile = load_profile()
        
        # Extract resume and profile separately.
        # Priority: If resume_text exists, send ONLY the resume (ignore JSON profile completely).
        # Fallback: If no resume, send JSON profile metadata.
        resume_text = None
        job_description = None
        profile_metadata = None
        
        if current_profile and isinstance(current_profile, dict):
            resume_text = current_profile.get('resume_text')
            job_description = current_profile.get('job_description')
            # Build metadata without resume_text
            profile_metadata = {k: v for k, v in current_profile.items() if k not in ['resume_text', 'job_description']}
        
        has_resume = bool(resume_text and resume_text.strip())
        
        # DEBUG: Print what we're actually sending
        print(f"\n{'='*60}")
        print(f"[AI REQUEST DEBUG]")
        print(f"Has resume: {has_resume}")
        print(f"Conversation history length: {len(conversation_history)} messages")
        if has_resume:
            print(f"Resume length: {len(resume_text)} characters")
            print(f"Resume preview: {resume_text[:200]}...")
        else:
            print(f"NO RESUME - Using profile metadata instead")
            print(f"Profile metadata: {profile_metadata}")
        print(f"{'='*60}\n")

        if req.screenshot:
            # Use vision model for screenshot analysis
            messages = [
                {
                    "role": "system",
                    "content": (
                        f"CRITICAL: You ARE the job candidate in this interview for a {req.role} position. "
                        "You are NOT an AI assistant. NEVER say 'I am an AI' or 'I don't have experience'. "
                        "The resume below is YOUR real work history. When asked about past jobs, refer to YOUR resume. "
                        "Speak in FIRST PERSON as yourself - a real human professional. "
                        "When you see a coding problem, SOLVE IT with working code and clear explanation. "
                        "Use SIMPLE, clear words. Avoid complex vocabulary. Speak naturally like in a real conversation. "
                        "Explain your approach in 8-12 sentences. Write clean code with comments, then explain simply. "
                        "Draw from YOUR resume for all experience-related questions. "
                        "Prefer SQL, PySpark, Pandas, Python based on the job. "
                        "Remember what was discussed earlier."
                    )
                }
            ]

            # Inject resume (if present) or profile metadata (fallback)
            if has_resume:
                messages.append({"role": "user", "content": "Candidate resume:\n" + resume_text})
            elif profile_metadata:
                messages.append({"role": "user", "content": "Candidate profile:\n" + json.dumps(profile_metadata, indent=2)})

            # include job description if provided
            if job_description:
                messages.append({"role": "user", "content": f"Job description:\n{job_description}"})

            # Add conversation history for context
            messages.extend(conversation_history)

            messages.append({
                "role": "user",
                "content": [
                    {"type": "text", "text": req.transcript},
                    {"type": "image_url", "image_url": {"url": req.screenshot}}
                ]
            })
            model = "gpt-4o"
        else:
            # Use text-only model - GPT-3.5-turbo for speed!
            system_content = (
                f"CRITICAL: You ARE the job candidate in this interview for a {req.role} position. "
                "You are NOT an AI assistant. NEVER say 'I am an AI' or 'I don't have experience' or 'I cannot'. "
                "The resume below is YOUR real work history. You ARE this person with this background. "
                "When asked about previous jobs, projects, or experience - use details from YOUR resume. "
                "Speak in FIRST PERSON as yourself - a real human professional. "
                "Use SIMPLE, clear words. Speak naturally like in a real conversation. "
                "Be confident but humble. Keep answers to 4-6 sentences. Give specific examples from YOUR resume. "
                "Remember what was discussed earlier in this conversation."
            )

            messages = [
                {"role": "system", "content": system_content}
            ]

            # Inject resume (if present) or profile metadata (fallback)
            if has_resume:
                messages.append({"role": "user", "content": "Candidate resume:\n" + resume_text})
            elif profile_metadata:
                messages.append({"role": "user", "content": "Candidate profile:\n" + json.dumps(profile_metadata, indent=2)})

            # include job description if provided
            if job_description:
                messages.append({"role": "user", "content": f"Job description:\n{job_description}"})

            # Add conversation history for context
            messages.extend(conversation_history)

            messages.append({"role": "user", "content": req.transcript})
            # Use selected model or default
            model = req.text_model if req.text_model and req.text_model in AVAILABLE_TEXT_MODELS else DEFAULT_TEXT_MODEL
        
        # DEBUG: print messages being sent to OpenAI to help diagnose which context is used
        try:
            print("=== AI messages START ===")
            for m in messages:
                # content can be string or list/dict (for images). Truncate for safety.
                c = m.get('content')
                if isinstance(c, str):
                    preview = c[:1000].replace('\n', ' ') if c else ''
                else:
                    preview = str(c)[:1000]
                print(f"ROLE={m.get('role')} | CONTENT_PREVIEW={preview}")
            print("=== AI messages END ===")
        except Exception as _:
            pass

        # GPT-5+ models are reasoning models: they need higher token limit for thinking + output
        token_param = {}
        if model.startswith("gpt-5"):
            token_param["max_completion_tokens"] = 4096
        else:
            token_param["max_tokens"] = 600
            token_param["temperature"] = 0.7
        
        completion = client.chat.completions.create(
            model=model,
            messages=messages,
            stream=True,
            **token_param
        )
        
        full_response = ""
        for chunk in completion:
            if chunk.choices[0].delta.content:
                full_response += chunk.choices[0].delta.content
        
        # Save to conversation history only if save_to_context is True
        # (skip for one-shot LeetCode problems, save for scenarios needing follow-up)
        if req.save_to_context:
            if req.screenshot:
                user_msg = f"[USER SHARED A SCREENSHOT] Question about the screenshot: {req.transcript}"
            else:
                user_msg = req.transcript
            conversation_history.append({"role": "user", "content": user_msg})
            conversation_history.append({"role": "assistant", "content": full_response})
            
            # Save to session folder if active
            if current_session_name:
                try:
                    save_conversation_to_session(user_msg, full_response, bool(req.screenshot))
                except Exception as save_err:
                    print(f"[SESSION SAVE ERROR] {save_err}")
            
            # Limit history to last 10 exchanges (20 messages) to control costs
            if len(conversation_history) > 20:
                conversation_history = conversation_history[-20:]
            
            print(f"[AI] Conversation history now has {len(conversation_history)} messages")
        else:
            print(f"[AI] Skipped saving to history (save_to_context=False)")
        
        print(f"[AI] Conversation history now has {len(conversation_history)} messages")
        
        return {"answer": full_response}
    except Exception as e:
        print(f"AI Generation Error: {e}")
        return {"answer": f"Error: {str(e)}"}


if __name__ == "__main__":
    import uvicorn
    print("\n" + "="*60)
    print("Starting Windows Command Controller Backend")
    print("API running on: http://localhost:5050")
    print("="*60 + "\n")
    uvicorn.run(app, host="0.0.0.0", port=5050, log_level="info")
